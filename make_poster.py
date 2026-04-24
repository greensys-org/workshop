import qrcode
import io
import os
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── QR code ──────────────────────────────────────────────────────────────────
qr = qrcode.QRCode(version=1, box_size=8, border=2)
qr.add_data('https://greensys-org.github.io/workshop/program/')
qr.make(fit=True)
qr_img = qr.make_image(fill_color="black", back_color="white")
qr_buffer = io.BytesIO()
qr_img.save(qr_buffer, format='PNG')
qr_buffer.seek(0)

# ── Circular crop helper ──────────────────────────────────────────────────────
def make_circle_image(path, size=200):
    img = Image.open(path).convert("RGBA")
    img = img.resize((size, size), Image.LANCZOS)
    mask = Image.new("L", (size, size), 0)
    draw = ImageDraw.Draw(mask)
    draw.ellipse((0, 0, size, size), fill=255)
    result = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    result.paste(img, mask=mask)
    # White background for Word compatibility
    bg = Image.new("RGB", (size, size), (255, 255, 255))
    bg.paste(result, mask=result.split()[3])
    buf = io.BytesIO()
    bg.save(buf, format='PNG')
    buf.seek(0)
    return buf

# ── Helpers to style table borders ───────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = OxmlElement(tag)
        element.set(qn('w:val'), kwargs.get(edge, 'none'))
        element.set(qn('w:sz'), '0')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Remove default paragraph spacing
style = doc.styles['Normal']
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(0)

# ── Banner / heading ──────────────────────────────────────────────────────────
# Top colour bar via a 1-cell table
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = banner.cell(0, 0)
set_cell_bg(banner_cell, '1F5C8B')  # dark blue
set_cell_border(banner_cell)
bp = banner_cell.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(8)
bp.paragraph_format.space_after  = Pt(4)
run = bp.add_run('GreenSys 2026')
run.font.size  = Pt(30)
run.font.bold  = True
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
run.font.name  = 'Calibri'

bp2 = banner_cell.add_paragraph()
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp2.paragraph_format.space_before = Pt(0)
bp2.paragraph_format.space_after  = Pt(8)
run2 = bp2.add_run('Keynote Speakers')
run2.font.size  = Pt(22)
run2.font.bold  = True
run2.font.color.rgb = RGBColor(0xA8, 0xD1, 0xF0)
run2.font.name  = 'Calibri'

# Sub-header
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(5)
sub.paragraph_format.space_after  = Pt(6)
sr = sub.add_run('Monday, 27 April 2026  ·  Edinburgh, UK  ·  Co-located with EuroSys 2026')
sr.font.size  = Pt(10)
sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sr.font.name  = 'Calibri'

# ── Speakers data ─────────────────────────────────────────────────────────────
IMG = '/Users/shashi/code/workshop/assets/img/'
speakers = [
    {
        'name':        'Loic Lannelongue',
        'affiliation': 'University of Cambridge, UK',
        'talk':        'Un/sustainable computing: what can we do about it?',
        'time':        '09:15 – 10:00',
        'photo':       IMG + 'loic_lannelongue.jpg',
    },
    {
        'name':        'Pierre Jacquet',
        'affiliation': 'Université du Québec (ETS Montréal) and OVHcloud',
        'talk':        'Monitoring GPUs in Cloud Platforms: Challenges and Opportunities for Orchestration',
        'time':        '11:00 – 11:45',
        'photo':       IMG + 'pierre_jacquet.jpg',
    },
    {
        'name':        'Edoardo M. Ponti',
        'affiliation': 'University of Edinburgh, UK',
        'talk':        'Adaptive Foundation Models for Efficient and Long-Horizon AI',
        'time':        '14:00 – 14:50',
        'photo':       IMG + 'edoardo_ponti.jpeg',
    },
    {
        'name':        'Maximilian Boether',
        'affiliation': 'ETH Zurich and DatologyAI',
        'talk':        'Efficient Data Mixing and Loading for Foundation Model Training',
        'time':        '15:35 – 16:20',
        'photo':       IMG + 'maximilian_boether.jpg',
    },
]

# ── Speaker cards in a 2×2 table ─────────────────────────────────────────────
card_bg_colors = ['EEF4F8', 'F4EEF8', 'EEF8F0', 'F8F4EE']

grid = doc.add_table(rows=2, cols=2)
grid.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set equal column widths
col_width = Cm(8.75)
for col in grid.columns:
    for cell in col.cells:
        cell.width = col_width

for i, sp in enumerate(speakers):
    row = i // 2
    col = i % 2
    cell = grid.cell(row, col)
    set_cell_bg(cell, card_bg_colors[i])
    set_cell_border(cell, top='single', left='single', bottom='single', right='single')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Photo
    pp = cell.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(10)
    pp.paragraph_format.space_after  = Pt(4)
    if os.path.exists(sp['photo']):
        img_buf = make_circle_image(sp['photo'], size=220)
        pp.add_run().add_picture(img_buf, width=Inches(1.4))

    # Name
    np = cell.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np.paragraph_format.space_before = Pt(4)
    np.paragraph_format.space_after  = Pt(2)
    nr = np.add_run(sp['name'])
    nr.font.bold  = True
    nr.font.size  = Pt(13)
    nr.font.name  = 'Calibri'
    nr.font.color.rgb = RGBColor(0x1F, 0x5C, 0x8B)

    # Affiliation
    ap = cell.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(6)
    ar = ap.add_run(sp['affiliation'])
    ar.font.size   = Pt(9)
    ar.font.italic = True
    ar.font.name   = 'Calibri'
    ar.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Talk title
    tp = cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(6)
    tr = tp.add_run(f'\u201c{sp["talk"]}\u201d')
    tr.font.size  = Pt(10)
    tr.font.name  = 'Calibri'
    tr.font.color.rgb = RGBColor(0xD9, 0x6B, 0x1F)

    # Time
    timep = cell.add_paragraph()
    timep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timep.paragraph_format.space_before = Pt(0)
    timep.paragraph_format.space_after  = Pt(10)
    timer = timep.add_run(f'\u23f0  {sp["time"]}')
    timer.font.size  = Pt(11)
    timer.font.bold  = True
    timer.font.name  = 'Calibri'
    timer.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

# ── QR code row ───────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(4)

qr_table = doc.add_table(rows=1, cols=2)
qr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

left_cell  = qr_table.cell(0, 0)
right_cell = qr_table.cell(0, 1)
set_cell_border(left_cell)
set_cell_border(right_cell)

# Left: website label
lp = left_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
lp.paragraph_format.space_before = Pt(6)
lr = lp.add_run('greensys-org.github.io/workshop/program/')
lr.font.size  = Pt(8)
lr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
lr.font.name  = 'Calibri'

# Right: QR code
rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_before = Pt(2)
rr = rp.add_run()
rr.add_picture(qr_buffer, width=Inches(0.9))

rp2 = right_cell.add_paragraph()
rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp2r = rp2.add_run('Scan for full program')
rp2r.font.size  = Pt(7)
rp2r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
rp2r.font.name  = 'Calibri'

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/shashi/code/workshop/GreenSys2026_Keynote_Poster.docx'
doc.save(out)
print(f'Saved: {out}')
